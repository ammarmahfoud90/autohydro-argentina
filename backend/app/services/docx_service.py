"""
Word (DOCX) Memoria de Cálculo Hidrológico generator using python-docx.

Produces a professional Argentine-standard hydrological calculation report
in Microsoft Word format, mirroring the structure and content of the PDF report.
"""

from io import BytesIO
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.map_image_service import generate_basin_map_image

# ── Brand colours (hex strings for OOXML, RGBColor for runs) ──────────────────

_NAVY_HEX = "1a365d"
_LIGHT_BLUE_HEX = "dbeafe"
_LIGHT_GRAY_HEX = "f3f4f6"

_NAVY = RGBColor(0x1A, 0x36, 0x5D)
_BLUE = RGBColor(0x25, 0x63, 0xEB)
_GRAY = RGBColor(0x6B, 0x72, 0x80)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_METHOD_NAMES: dict[str, str] = {
    "rational": "Método Racional",
    "modified_rational": "Método Racional Modificado",
    "scs_cn": "Método SCS-CN (Soil Conservation Service)",
}

_INFRASTRUCTURE_NAMES: dict[str, str] = {
    "alcantarilla_menor": "Alcantarilla menor (< 1 m)",
    "alcantarilla_mayor": "Alcantarilla mayor (1–3 m)",
    "puente_menor": "Puente menor (3–10 m luz)",
    "puente_mayor": "Puente mayor (> 10 m luz)",
    "canal_urbano": "Canal urbano",
    "canal_rural": "Canal rural",
    "defensa_costera": "Defensa costera",
}

_RISK_LABELS_ES: dict[str, str] = {
    "muy_bajo": "MUY BAJO",
    "bajo": "BAJO",
    "moderado": "MODERADO",
    "alto": "ALTO",
    "muy_alto": "MUY ALTO",
}


# ── OOXML helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Set table cell background colour via OOXML shading element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_field(run: Any) -> None:
    """Inject a PAGE field into a run element."""
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


# ── Generator class ────────────────────────────────────────────────────────────

class MemoriaCalculoDocxGenerator:
    """Generate a professional hydrological calculation report as .docx."""

    def __init__(
        self,
        project_name: str = "Proyecto Hidrológico",
        location: str = "Argentina",
        engineer_name: str = "Ing. Ammar Mahfoud",
        client: Optional[str] = None,
        language: str = "es",
    ) -> None:
        self.project_name = project_name
        self.location = location
        self.engineer_name = engineer_name
        self.client = client
        self.language = language

    # ── Document initialisation ────────────────────────────────────────────────

    def _setup_document(self) -> Document:
        doc = Document()

        # A4 page size and 2.5 cm margins
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.2)

        self._setup_footer(section)

        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        return doc

    def _setup_footer(self, section: Any) -> None:
        """Add footer: credit text (left) + page number (right)."""
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()

        # Split footer with a centre tab + right tab
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run_left = para.add_run(
            "Generado con AutoHydro Argentina \u2014 Desarrollado por Ing. Ammar Mahfoud"
        )
        run_left.font.size = Pt(7.5)
        run_left.font.color.rgb = _GRAY

        run_tab = para.add_run("\t\tP\u00e1gina\u00a0")
        run_tab.font.size = Pt(7.5)
        run_tab.font.color.rgb = _GRAY

        run_pg = para.add_run()
        run_pg.font.size = Pt(7.5)
        run_pg.font.color.rgb = _GRAY
        _add_page_number_field(run_pg)

    # ── Typography helpers ─────────────────────────────────────────────────────

    def _add_paragraph_border(
        self,
        para: Any,
        side: str = "bottom",
        hex_color: str = _NAVY_HEX,
        sz: int = 6,
    ) -> None:
        """Add a top or bottom border to a paragraph (simulates HR rules)."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        border_el = OxmlElement(f"w:{side}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), str(sz))
        border_el.set(qn("w:space"), "1")
        border_el.set(qn("w:color"), hex_color)
        pBdr.append(border_el)
        pPr.append(pBdr)

    def _h1(self, doc: Document, text: str) -> None:
        """Level-1 section heading with a light-blue bottom rule."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = _NAVY
        self._add_paragraph_border(para, "bottom", _LIGHT_BLUE_HEX, sz=6)

    def _h2(self, doc: Document, text: str) -> None:
        """Level-2 sub-heading."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = _BLUE

    def _body(self, doc: Document, text: str, justify: bool = True) -> None:
        """Justified (or left-aligned) body paragraph."""
        para = doc.add_paragraph(text)
        para.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        )
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(10)

    # ── Table helper ───────────────────────────────────────────────────────────

    def _make_table(
        self,
        doc: Document,
        data: list[list],
        col_widths: Optional[list] = None,
        highlight_last: bool = False,
    ) -> Any:
        """
        Build a styled data table with a navy header row and alternating body rows.
        data[0] is the header row.
        """
        if not data:
            return None

        n_rows = len(data)
        n_cols = len(data[0])

        if col_widths is None:
            if n_cols == 2:
                col_widths = [Cm(7), Cm(9)]
            elif n_cols == 3:
                col_widths = [Cm(7), Cm(4), Cm(5)]
            else:
                width_each = Cm(16) / n_cols
                col_widths = [width_each] * n_cols

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.width = col_widths[col_idx] if col_idx < len(col_widths) else Cm(4)

                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(str(cell_text))
                run.font.size = Pt(9)

                if row_idx == 0:
                    # Header row: navy background, white bold centred text
                    _set_cell_bg(cell, _NAVY_HEX)
                    run.font.color.rgb = _WHITE
                    run.font.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # First column: bold navy
                    if col_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = _NAVY
                    # Alternating row background
                    if row_idx % 2 == 0:
                        _set_cell_bg(cell, _LIGHT_GRAY_HEX)
                    # Highlight last row (result row)
                    if highlight_last and row_idx == n_rows - 1:
                        _set_cell_bg(cell, _LIGHT_BLUE_HEX)
                        run.font.bold = True

        doc.add_paragraph()  # visual spacing after table
        return table

    # ── Cover page ─────────────────────────────────────────────────────────────

    def _build_cover(self, doc: Document) -> None:
        # Vertical spacer
        for _ in range(3):
            doc.add_paragraph()

        # Top rule (empty paragraph with thick bottom border)
        rule_top = doc.add_paragraph()
        rule_top.paragraph_format.space_after = Pt(14)
        self._add_paragraph_border(rule_top, "bottom", _NAVY_HEX, sz=18)

        # Main title
        para_title = doc.add_paragraph()
        para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_title.paragraph_format.space_after = Pt(4)
        run_title = para_title.add_run("MEMORIA DE C\u00c1LCULO HIDROL\u00d3GICO")
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = _NAVY

        # Subtitle
        para_sub = doc.add_paragraph()
        para_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_sub.paragraph_format.space_after = Pt(4)
        run_sub = para_sub.add_run("An\u00e1lisis de Caudales de Dise\u00f1o")
        run_sub.font.size = Pt(13)
        run_sub.font.color.rgb = _GRAY

        # Thin blue separator rule
        rule_mid = doc.add_paragraph()
        rule_mid.paragraph_format.space_after = Pt(20)
        self._add_paragraph_border(rule_mid, "bottom", _LIGHT_BLUE_HEX, sz=4)

        doc.add_paragraph()

        # Project & location
        para_proj = doc.add_paragraph()
        para_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_proj.paragraph_format.space_after = Pt(4)
        r_pl = para_proj.add_run("Proyecto: ")
        r_pl.font.bold = True
        r_pl.font.size = Pt(12)
        r_pv = para_proj.add_run(self.project_name)
        r_pv.font.size = Pt(12)

        para_loc = doc.add_paragraph()
        para_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_loc.paragraph_format.space_after = Pt(4)
        r_ll = para_loc.add_run("Ubicaci\u00f3n: ")
        r_ll.font.bold = True
        r_ll.font.size = Pt(12)
        r_lv = para_loc.add_run(self.location)
        r_lv.font.size = Pt(12)

        for _ in range(4):
            doc.add_paragraph()

        # Info table (no borders — clean layout)
        info_rows = []
        if self.client:
            info_rows.append(("Comitente:", self.client))
        info_rows.append(("Profesional Responsable:", self.engineer_name))
        info_rows.append(("Fecha de emisi\u00f3n:", datetime.now().strftime("%d/%m/%Y")))
        info_rows.append(("Herramienta:", "AutoHydro Argentina v1.0"))

        tbl = doc.add_table(rows=len(info_rows), cols=2)
        for row_idx, (label, value) in enumerate(info_rows):
            row = tbl.rows[row_idx]
            row.cells[0].width = Cm(7)
            row.cells[1].width = Cm(9)

            row.cells[0].paragraphs[0].clear()
            r_lbl = row.cells[0].paragraphs[0].add_run(label)
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(10)
            r_lbl.font.color.rgb = _NAVY
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            row.cells[1].paragraphs[0].clear()
            r_val = row.cells[1].paragraphs[0].add_run(value)
            r_val.font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        # Bottom thick rule
        rule_bot = doc.add_paragraph()
        rule_bot.paragraph_format.space_after = Pt(14)
        self._add_paragraph_border(rule_bot, "bottom", _NAVY_HEX, sz=18)

        # Disclaimer
        disclaimer = (
            "AVISO: Esta memoria de c\u00e1lculo ha sido generada con AutoHydro Argentina. "
            "Los coeficientes IDF utilizados son de car\u00e1cter indicativo basados en la "
            "regionalizaci\u00f3n de Caama\u00f1o Nelli et al. (1999) e INA. Para dise\u00f1os "
            "definitivos, verifique con los estudios hidrol\u00f3gicos locales m\u00e1s "
            "recientes y la normativa vigente aplicable."
        )
        para_disc = doc.add_paragraph(disclaimer)
        para_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para_disc.runs:
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = _GRAY

        doc.add_page_break()

    # ── Table of contents ──────────────────────────────────────────────────────

    def _build_toc(self, doc: Document) -> None:
        self._h1(doc, "\u00cdNDICE")

        toc_items = [
            ("1.", "Objeto del Estudio"),
            ("2.", "Descripción de la Cuenca"),
            ("3.", "Ubicación de la Cuenca"),
            ("4.", "Análisis Pluviométrico — Curvas IDF"),
            ("5.", "Tiempo de Concentración"),
            ("6.", "Metodología de Cálculo"),
            ("7.", "Cálculos y Resultados"),
            ("8.", "Análisis e Interpretación"),
            ("9.", "Conclusiones y Recomendaciones"),
            ("Anexo A.", "Planilla de Cálculo Detallada"),
        ]

        for num, title in toc_items:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            r1 = para.add_run(f"{num} ")
            r1.font.bold = True
            r1.font.size = Pt(10)
            r2 = para.add_run(title)
            r2.font.size = Pt(10)

        doc.add_page_break()

    # ── Section 1: Objeto ──────────────────────────────────────────────────────

    def _build_section_objeto(self, doc: Document, ai_sections: dict) -> None:
        self._h1(doc, "1. OBJETO DEL ESTUDIO")
        text = ai_sections.get("objeto") or (
            "El presente estudio tiene por objeto determinar el caudal m\u00e1ximo de "
            "dise\u00f1o para la infraestructura hidrol\u00f3gica propuesta, mediante la "
            "aplicaci\u00f3n de metodolog\u00eda hidrol\u00f3gica apropiada para las "
            "caracter\u00edsticas de la cuenca de aporte."
        )
        self._body(doc, text)

    # ── Section 2: Cuenca ──────────────────────────────────────────────────────

    def _build_section_cuenca(
        self, doc: Document, data: dict, ai_sections: dict
    ) -> None:
        self._h1(doc, "2. DESCRIPCI\u00d3N DE LA CUENCA")

        text = ai_sections.get("descripcion_cuenca") or ""
        if text:
            self._body(doc, text)

        self._h2(doc, "Par\u00e1metros morfom\u00e9tricos:")

        morph = [
            ["Par\u00e1metro", "Valor", "Unidad"],
            ["\u00c1rea de la cuenca (A)", f"{data['area_km2']:.3f}", "km\u00b2"],
            ["Longitud del cauce principal (L)", str(data.get("length_km", "\u2014")), "km"],
            ["Pendiente media del cauce (S)", str(data.get("slope", "\u2014")), "m/m"],
        ]
        if data.get("elevation_diff_m"):
            morph.append(["Desnivel total (H)", f"{data['elevation_diff_m']:.1f}", "m"])
        if data.get("avg_elevation_m"):
            morph.append(["Altura media sobre cierre (Hm)", f"{data['avg_elevation_m']:.1f}", "m"])

        self._make_table(doc, morph)

    # ── Section 3: Ubicación de la cuenca ─────────────────────────────────────

    def _build_section_ubicacion(
        self, doc: Document, basin_polygon: Optional[list[list[float]]]
    ) -> None:
        self._h1(doc, "3. UBICACIÓN DE LA CUENCA")

        if basin_polygon and len(basin_polygon) >= 3:
            img_bytes = generate_basin_map_image(basin_polygon, width=800, height=500)
            if img_bytes:
                img_buf = BytesIO(img_bytes)
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run()
                run.add_picture(img_buf, width=Cm(14))
                caption = doc.add_paragraph("Figura 1: Delimitación de la cuenca de estudio")
                caption.paragraph_format.space_after = Pt(6)
                for run in caption.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = _GRAY
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self._body(
                    doc,
                    "No fue posible generar la imagen del mapa. "
                    "Verificar la conexión a internet para cargar las teselas base.",
                )
        else:
            self._body(
                doc,
                "No se proporcionó delimitación cartográfica de la cuenca. "
                "La ubicación fue definida mediante parámetros morfométricos ingresados manualmente.",
            )

    # ── Section 4: IDF ────────────────────────────────────────────────────────

    def _build_section_idf(self, doc: Document, data: dict) -> None:
        self._h1(doc, "4. ANÁLISIS PLUVIOMÉTRICO — CURVAS IDF")

        intro = (
            f"Los datos pluviom\u00e9tricos se obtuvieron de la estaci\u00f3n de referencia "
            f"{data['city']} (provincia de {data.get('province', '')})."
        )
        self._body(doc, intro)

        idf_verified = data.get("idf_verified", True)
        idf_source = data.get("idf_source", "—") if idf_verified else "Estimación regional (requiere validación local)"

        idf_data = [
            ["Parámetro", "Valor"],
            ["Estación de referencia", data["city"]],
            ["Fuente / Bibliografía", idf_source],
            ["Período de retorno (T)", f"{data['return_period']} años"],
            ["Duración de tormenta (t)", f"{data['duration_min']} minutos"],
            ["Intensidad de diseño (i)", f"{data['intensity_mm_hr']:.2f} mm/hr"],
        ]
        self._make_table(doc, idf_data)

        if not idf_verified:
            self._body(
                doc,
                "⚠ AVISO: Los datos IDF para esta localidad son estimaciones regionales "
                "interpoladas y no han sido validados contra registros pluviográficos locales. "
                "Para diseños definitivos, verificar con registros locales del SMN/INA.",
            )

        self._body(doc, "Fórmula IDF aplicada:  i = (a × T^b) / (t + c)^d", justify=False)
        self._body(
            doc,
            "donde i [mm/hr], T [años], t [minutos] y a, b, c, d son coeficientes regionales.",
        )

    # ── Section 4: Tc ─────────────────────────────────────────────────────────

    def _build_section_tc(self, doc: Document, data: dict) -> None:
        self._h1(doc, "5. TIEMPO DE CONCENTRACIÓN (Tc)")

        n = len(data.get("tc_results", []))
        self._body(
            doc,
            f"Se calcul\u00f3 el tiempo de concentraci\u00f3n mediante {n} f\u00f3rmula(s). "
            "El valor adoptado corresponde al promedio de los resultados obtenidos.",
        )

        tc_data = [["F\u00f3rmula", "Tc (hr)", "Tc (min)", "Aplicabilidad"]]
        for r in data.get("tc_results", []):
            appl = r["applicability"]
            if len(appl) > 55:
                appl = appl[:55] + "\u2026"
            tc_data.append([
                r["formulaName"],
                f"{r['tcHours']:.3f}",
                f"{r['tcMinutes']:.1f}",
                appl,
            ])
        tc_data.append([
            "TC ADOPTADO (promedio)",
            f"{data['tc_adopted_hours']:.3f}",
            f"{data['tc_adopted_minutes']:.1f}",
            "Valor de dise\u00f1o",
        ])

        self._make_table(
            doc,
            tc_data,
            col_widths=[Cm(4.5), Cm(2), Cm(2), Cm(7.5)],
            highlight_last=True,
        )

    # ── Section 5: Metodología ────────────────────────────────────────────────

    def _build_section_metodologia(
        self, doc: Document, data: dict, ai_sections: dict
    ) -> None:
        self._h1(doc, "6. METODOLOGÍA DE CÁLCULO")

        method_name = _METHOD_NAMES.get(data.get("method", ""), data.get("method", ""))
        para = doc.add_paragraph()
        r1 = para.add_run("M\u00e9todo seleccionado: ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = para.add_run(method_name)
        r2.font.size = Pt(10)

        text = ai_sections.get("metodologia") or ""
        if text:
            self._body(doc, text)

        self._h2(doc, "F\u00f3rmula aplicada:")
        method = data.get("method", "")
        if method == "rational":
            self._body(doc, "Q = C \u00d7 i \u00d7 A / 3.6", justify=False)
            self._body(doc, "donde Q [m\u00b3/s], C [-], i [mm/hr], A [km\u00b2]")
        elif method == "modified_rational":
            self._body(doc, "Q = C \u00d7 i \u00d7 A \u00d7 K / 3.6", justify=False)
            self._body(doc, "K = 1 \u2212 (A^0.1 \u2212 1) / 7  (T\u00e9mez)", justify=False)
        elif method == "scs_cn":
            self._body(
                doc,
                "S = 25400/CN \u2212 254  |  Ia = \u03bb \u00d7 S  |  Q = (P\u2212Ia)\u00b2 / (P\u2212Ia+S)",
                justify=False,
            )
            self._body(
                doc, "Qp = 0.208 \u00d7 A \u00d7 Q / Tp  |  Tp = 0.6 \u00d7 Tc", justify=False
            )
            lam = (
                "0.05 (Pampa H\u00fameda)"
                if data.get("use_pampa_lambda")
                else "0.20 (est\u00e1ndar)"
            )
            self._body(doc, f"Abstracci\u00f3n inicial: \u03bb = {lam}")

    # ── Section 6: Cálculos ───────────────────────────────────────────────────

    def _build_section_calculos(self, doc: Document, data: dict) -> None:
        self._h1(doc, "7. CÁLCULOS Y RESULTADOS")
        self._h2(doc, "7.1 Parámetros de cálculo")

        param_rows: list[list] = [["Par\u00e1metro", "Valor", "Unidad"]]
        if data.get("runoff_coeff") is not None:
            param_rows.append(
                ["Coeficiente de escorrent\u00eda (C)", f"{data['runoff_coeff']:.2f}", "\u2014"]
            )
        if data.get("cn") is not None:
            param_rows += [
                ["N\u00famero de Curva compuesto (CN)", f"{data['cn']:.1f}", "\u2014"],
                ["Retenci\u00f3n m\u00e1xima potencial (S)", f"{data['s_mm']:.1f}", "mm"],
                ["Abstracci\u00f3n inicial (Ia)", f"{data['ia_mm']:.1f}", "mm"],
                ["L\u00e1mina de escorrent\u00eda directa (Q)", f"{data['runoff_depth_mm']:.1f}", "mm"],
            ]
        if data.get("areal_reduction_k") is not None:
            param_rows.append(
                ["Factor de reducci\u00f3n areal (K)", f"{data['areal_reduction_k']:.4f}", "\u2014"]
            )
        param_rows += [
            ["Intensidad de dise\u00f1o (i)", f"{data['intensity_mm_hr']:.2f}", "mm/hr"],
            ["Tiempo de concentraci\u00f3n adoptado (Tc)", f"{data['tc_adopted_hours']:.3f}", "hr"],
            ["\u00c1rea de la cuenca (A)", f"{data['area_km2']:.3f}", "km\u00b2"],
        ]
        self._make_table(doc, param_rows)

        self._h2(doc, "7.2 Caudal de diseño")
        result_rows = [
            ["Par\u00e1metro", "Valor", "Unidad"],
            ["Caudal pico de dise\u00f1o (Q)", f"{data['peak_flow_m3s']:.3f}", "m\u00b3/s"],
            ["Caudal espec\u00edfico (q)", f"{data['specific_flow_m3s_km2']:.3f}", "m\u00b3/s/km\u00b2"],
        ]
        self._make_table(doc, result_rows, highlight_last=True)

        if data.get("method_comparison"):
            self._h2(doc, "7.3 Comparación de métodos")
            comp_rows = [["M\u00e9todo", "Q (m\u00b3/s)", "Tc (hr)", "Observaciones"]]
            for r in data["method_comparison"]:
                notes = r["notes"]
                if len(notes) > 50:
                    notes = notes[:50] + "\u2026"
                comp_rows.append([
                    r["methodName"],
                    f"{r['peakFlow']:.3f}",
                    f"{r['tc']:.3f}",
                    notes,
                ])
            self._make_table(doc, comp_rows, col_widths=[Cm(5), Cm(2.5), Cm(2.5), Cm(6)])

    # ── Section 6b: CN Sensitivity ────────────────────────────────────────────

    def _build_section_sensibilidad(self, doc: Document, data: dict) -> None:
        self._h2(doc, "7.4 Análisis de Sensibilidad del CN")

        self._body(
            doc,
            "Este análisis muestra cómo varía el caudal pico ante cambios de ±5 unidades "
            "en el Número de Curva. El CN es el parámetro con mayor incertidumbre en el "
            "método SCS-CN; variaciones de ±5 unidades son habituales en la práctica y "
            "pueden producir diferencias significativas en el caudal de diseño.",
        )

        sensitivity = data.get("cn_sensitivity") or []
        if not sensitivity:
            return

        rows: list[list] = [["Número de Curva", "Q pico (m³/s)", "Variación"]]
        for s in sensitivity:
            label_cn = f"{s['label']} ({s['cn']:.0f})"
            q_str = f"{s['peak_flow_m3s']:.3f}"
            if s["label"] == "CN":
                var_str = "Base"
            else:
                sign = "+" if s["variation_pct"] > 0 else ""
                var_str = f"{sign}{s['variation_pct']:.1f}%"
            rows.append([label_cn, q_str, var_str])

        # Highlight base row (index 2, which is the "CN" row)
        tbl = self._make_table(doc, rows, col_widths=[Cm(6), Cm(4), Cm(4)])
        # Override base row color (index 2) after table is built
        if tbl is not None and len(tbl.rows) > 2:
            base_row = tbl.rows[2]
            for cell in base_row.cells:
                _set_cell_bg(cell, _LIGHT_BLUE_HEX)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True

        note = doc.add_paragraph(
            "NOTA: Se recomienda adoptar el CN base como valor de diseño y usar CN+5 "
            "como escenario conservador de verificación."
        )
        note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in note.runs:
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = _GRAY

    # ── Section 7: Análisis ───────────────────────────────────────────────────

    def _build_section_analisis(
        self,
        doc: Document,
        data: dict,
        ai_interpretation: str,
        ai_sections: dict,
    ) -> None:
        self._h1(doc, "8. ANÁLISIS E INTERPRETACIÓN")

        risk = data.get("risk_level", "moderado")
        risk_label = _RISK_LABELS_ES.get(risk, risk.upper())
        infra = _INFRASTRUCTURE_NAMES.get(data.get("infrastructure_type", ""), "\u2014")

        para = doc.add_paragraph()
        r1 = para.add_run("Nivel de riesgo hidrol\u00f3gico: ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = para.add_run(f"{risk_label} \u2014 {infra}")
        r2.font.size = Pt(10)

        rec = data.get("risk_recommendations") or {}
        if rec:
            rec_rows: list[list] = [["Aspecto", "Descripci\u00f3n"]]
            if rec.get("general"):
                rec_rows.append(["Situaci\u00f3n general", rec["general"]])
            if rec.get("action"):
                rec_rows.append(["Acci\u00f3n recomendada", rec["action"]])
            if rec.get("verification"):
                rec_rows.append(["Verificaciones requeridas", rec["verification"]])
            if rec.get("period_note"):
                rec_rows.append(["Nota sobre per\u00edodo de retorno", rec["period_note"]])
            self._make_table(doc, rec_rows, col_widths=[Cm(5), Cm(11)])

        if ai_sections.get("analisis_resultados"):
            self._body(doc, ai_sections["analisis_resultados"])

        if ai_interpretation:
            self._h2(doc, "Interpretaci\u00f3n t\u00e9cnica asistida por IA:")
            for para_text in ai_interpretation.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    self._body(doc, para_text)

    # ── Section 8: Conclusiones ───────────────────────────────────────────────

    def _build_section_conclusiones(
        self, doc: Document, data: dict, ai_sections: dict
    ) -> None:
        self._h1(doc, "9. CONCLUSIONES Y RECOMENDACIONES")

        text = ai_sections.get("conclusiones") or (
            f"Se determin\u00f3 un caudal pico de dise\u00f1o de "
            f"{data.get('peak_flow_m3s', 0):.3f}\u00a0m\u00b3/s para un per\u00edodo de "
            f"retorno de {data.get('return_period', '\u2014')} a\u00f1os."
        )
        self._body(doc, text)

        summary_rows = [
            ["Par\u00e1metro", "Valor adoptado"],
            ["Caudal de dise\u00f1o", f"Q = {data.get('peak_flow_m3s', 0):.3f} m\u00b3/s"],
            ["Per\u00edodo de retorno", f"T = {data.get('return_period', '\u2014')} a\u00f1os"],
            ["M\u00e9todo de c\u00e1lculo", _METHOD_NAMES.get(data.get("method", ""), "\u2014")],
            ["Nivel de riesgo", _RISK_LABELS_ES.get(data.get("risk_level", ""), "\u2014")],
        ]
        self._make_table(doc, summary_rows)

        disclaimer = (
            "LIMITACIONES: Los resultados de esta memoria son estimaciones para etapas de "
            "anteproyecto basadas en la regionalizaci\u00f3n de Caama\u00f1o Nelli et al. (1999). "
            "Para proyectos definitivos se recomienda verificar con pluviograf\u00eda local "
            "actualizada, estudios de campo y la normativa municipal o provincial vigente."
        )
        para_disc = doc.add_paragraph(disclaimer)
        para_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para_disc.runs:
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = _GRAY

    # ── Annex: Detailed calculation sheet ─────────────────────────────────────

    def _build_annex(self, doc: Document, data: dict) -> None:
        self._h1(doc, "ANEXO A \u2014 PLANILLA DE C\u00c1LCULO DETALLADA")

        flat: list[tuple[str, str]] = [
            ("Ciudad de referencia", data.get("city", "\u2014")),
            ("Provincia", data.get("province", "\u2014")),
            ("Per\u00edodo de retorno (T)", f"{data.get('return_period', '\u2014')} a\u00f1os"),
            ("Duraci\u00f3n de tormenta (t)", f"{data.get('duration_min', '\u2014')} min"),
            ("\u00c1rea de la cuenca (A)", f"{data.get('area_km2', '\u2014')} km\u00b2"),
            ("Longitud del cauce (L)", f"{data.get('length_km', '\u2014')} km"),
            ("Pendiente media (S)", f"{data.get('slope', '\u2014')} m/m"),
            ("Intensidad IDF (i)", f"{data.get('intensity_mm_hr', '\u2014')} mm/hr"),
            ("Fuente IDF", data.get("idf_source", "\u2014")),
            (
                "Tc adoptado",
                f"{data.get('tc_adopted_hours', '\u2014')} hr  /  "
                f"{data.get('tc_adopted_minutes', '\u2014')} min",
            ),
            ("M\u00e9todo", _METHOD_NAMES.get(data.get("method", ""), "\u2014")),
        ]

        if data.get("runoff_coeff") is not None:
            flat.append(("Coeficiente C", f"{data['runoff_coeff']:.3f}"))
        if data.get("cn") is not None:
            flat += [
                ("N\u00famero de Curva (CN)", f"{data['cn']:.1f}"),
                ("Retenci\u00f3n m\u00e1x. (S)", f"{data['s_mm']:.2f} mm"),
                ("Abstracci\u00f3n inicial (Ia)", f"{data['ia_mm']:.2f} mm"),
                ("L\u00e1mina escorrent\u00eda (Q)", f"{data['runoff_depth_mm']:.2f} mm"),
            ]
        if data.get("areal_reduction_k") is not None:
            flat.append(("Factor areal (K)", f"{data['areal_reduction_k']:.4f}"))

        flat += [
            ("Caudal pico (Q)", f"{data.get('peak_flow_m3s', '\u2014')} m\u00b3/s"),
            ("Caudal espec\u00edfico (q)", f"{data.get('specific_flow_m3s_km2', '\u2014')} m\u00b3/s/km\u00b2"),
            ("Nivel de riesgo", _RISK_LABELS_ES.get(data.get("risk_level", ""), "\u2014")),
            (
                "Infraestructura",
                _INFRASTRUCTURE_NAMES.get(data.get("infrastructure_type", ""), "\u2014"),
            ),
            ("Fecha de c\u00e1lculo", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ]

        all_rows = [["Campo", "Valor"]] + [[k, str(v)] for k, v in flat]
        self._make_table(doc, all_rows)

    # ── Main generate ──────────────────────────────────────────────────────────

    def generate(
        self,
        calculation_data: dict[str, Any],
        ai_interpretation: str = "",
        ai_recommendations: Any = "",
        basin_polygon: Optional[list[list[float]]] = None,
    ) -> BytesIO:
        """
        Generate the complete Word report.

        Args:
            calculation_data: Full CalculationResponse dict
            ai_interpretation: AI-generated interpretation text
            ai_recommendations: AI-generated report sections dict or empty string

        Returns:
            BytesIO buffer containing the .docx file
        """
        ai_sections: dict[str, str] = (
            ai_recommendations if isinstance(ai_recommendations, dict) else {}
        )

        doc = self._setup_document()

        # Cover + TOC
        self._build_cover(doc)
        self._build_toc(doc)

        # Main sections (mirror PDF section order)
        self._build_section_objeto(doc, ai_sections)
        self._build_section_cuenca(doc, calculation_data, ai_sections)
        self._build_section_ubicacion(doc, basin_polygon)
        self._build_section_idf(doc, calculation_data)
        self._build_section_tc(doc, calculation_data)
        self._build_section_metodologia(doc, calculation_data, ai_sections)
        doc.add_page_break()
        self._build_section_calculos(doc, calculation_data)
        if calculation_data.get("cn_sensitivity"):
            self._build_section_sensibilidad(doc, calculation_data)
        self._build_section_analisis(doc, calculation_data, ai_interpretation, ai_sections)
        self._build_section_conclusiones(doc, calculation_data, ai_sections)
        doc.add_page_break()

        # Annex
        self._build_annex(doc, calculation_data)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
